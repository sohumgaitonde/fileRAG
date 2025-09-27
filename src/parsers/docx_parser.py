"""
DOCX file parser using python-docx.

Extracts text content from Microsoft Word documents while preserving formatting.
"""

import os
from typing import Dict, List
from docx import Document
from docx.document import Document as DocumentType


class DOCXParser:
    """Parser for DOCX files."""
    
    def __init__(self):
        pass
    
    def parse(self, file_path: str) -> dict:
        """Parse DOCX file and extract text content."""
        try:
            # Load the document
            doc = Document(file_path)
            
            # Extract text content
            content = self._extract_text(doc)
            
            # Extract metadata
            metadata = self.extract_metadata(file_path)
            metadata.update({
                "file_path": file_path,
                "file_type": "docx",
                "paragraphs": len(doc.paragraphs),
                "content_length": len(content)
            })
            
            return {
                "content": content,
                "metadata": metadata
            }
            
        except Exception as e:
            return {
                "content": "",
                "metadata": {
                    "file_path": file_path,
                    "file_type": "docx",
                    "error": str(e),
                    "paragraphs": 0,
                    "title": "",
                    "author": ""
                }
            }
    
    def _extract_text(self, doc: DocumentType) -> str:
        """Extract text content from document paragraphs."""
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Skip empty paragraphs
                text_parts.append(paragraph.text.strip())
        
        return "\n\n".join(text_parts)
    
    def extract_metadata(self, file_path: str) -> dict:
        """Extract metadata from DOCX file."""
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            # Get file stats
            file_stats = os.stat(file_path)
            
            return {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": core_props.created.isoformat() if core_props.created else "",
                "modified": core_props.modified.isoformat() if core_props.modified else "",
                "file_size": file_stats.st_size,
                "tables": len(doc.tables),
                "images": len(doc.inline_shapes)
            }
            
        except Exception as e:
            return {
                "title": "",
                "author": "",
                "subject": "",
                "keywords": "",
                "created": "",
                "modified": "",
                "file_size": 0,
                "tables": 0,
                "images": 0,
                "error": str(e)
            }
